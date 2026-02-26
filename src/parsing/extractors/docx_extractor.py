"""DOCX document extractor using python-docx library."""

import logging
from typing import Any, Dict, BinaryIO, List

from src.parsing.extractors.base import BaseExtractor
from src.parsing.exceptions import (
    CorruptedFileException,
    ExtractionFailedException,
)

logger = logging.getLogger(__name__)


class DOCXExtractor(BaseExtractor):
    """DOCX extractor using python-docx library.

    Extracts text content, formatting information, tables, lists,
    and metadata from Microsoft Word DOCX documents.
    """

    def extract_text(self, file_obj: BinaryIO) -> str:
        """Extract text content from DOCX file.

        Args:
            file_obj: File object opened in binary mode

        Returns:
            str: Extracted text content

        Raises:
            CorruptedFileException: If DOCX is corrupted
            ExtractionFailedException: If text extraction fails
        """
        logger.info(f"Extracting text from DOCX: {self.filename}")

        try:
            import docx
        except ImportError:
            logger.error("python-docx library not installed")
            raise ExtractionFailedException(
                self.filename, "text", "python-docx library not available"
            )

        try:
            file_obj.seek(0)
            document = docx.Document(file_obj)

            text_parts: List[str] = []

            # Extract text from paragraphs
            for para in document.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            text = "\n".join(text_parts)

            logger.info(
                f"Successfully extracted {len(text)} characters from {self.filename}"
            )

            return text

        except Exception as e:
            logger.error(
                f"Failed to extract text from {self.filename}: {str(e)}",
                exc_info=True,
            )
            if "not a zip file" in str(e).lower() or "corrupt" in str(e).lower():
                raise CorruptedFileException(
                    self.filename, f"Invalid DOCX format: {str(e)}"
                )
            raise ExtractionFailedException(self.filename, "text", str(e))

    def extract_structured_data(self, file_obj: BinaryIO) -> Dict[str, Any]:
        """Extract structured metadata and formatting from DOCX.

        Args:
            file_obj: File object opened in binary mode

        Returns:
            Dict containing metadata, tables, lists, and formatting info

        Raises:
            ExtractionFailedException: If structured data extraction fails
        """
        logger.info(f"Extracting structured data from DOCX: {self.filename}")

        structured_data: Dict[str, Any] = {
            "metadata": {},
            "tables": [],
            "lists": [],
            "formatting": {
                "paragraph_count": 0,
                "has_bold": False,
                "has_italic": False,
                "has_underline": False,
            },
        }

        try:
            import docx

            file_obj.seek(0)
            document = docx.Document(file_obj)

            # Extract core properties metadata
            structured_data["metadata"] = self._extract_metadata(document)

            # Extract tables with structure
            structured_data["tables"] = self._extract_tables(document)

            # Extract lists
            structured_data["lists"] = self._extract_lists(document)

            # Extract formatting information
            structured_data["formatting"] = self._extract_formatting(document)

            logger.info(
                f"Extracted structured data from {self.filename}: "
                f"{structured_data['formatting']['paragraph_count']} paragraphs, "
                f"{len(structured_data['tables'])} tables"
            )

            return structured_data

        except Exception as e:
            logger.error(
                f"Failed to extract structured data from {self.filename}: {str(e)}",
                exc_info=True,
            )
            raise ExtractionFailedException(
                self.filename, "structured_data", str(e)
            )

    def _extract_metadata(self, document: Any) -> Dict[str, Any]:
        """Extract document metadata from core properties.

        Args:
            document: python-docx Document object

        Returns:
            Dict containing document metadata
        """
        metadata: Dict[str, Any] = {}

        try:
            core_props = document.core_properties

            # Extract common metadata fields
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.comments:
                metadata["comments"] = core_props.comments
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by

        except Exception as e:
            logger.warning(f"Failed to extract metadata from {self.filename}: {str(e)}")  # noqa: E501

        return metadata

    def _extract_tables(self, document: Any) -> List[Dict[str, Any]]:
        """Extract tables with structure from document.

        Args:
            document: python-docx Document object

        Returns:
            List of dictionaries containing table data
        """
        tables_data: List[Dict[str, Any]] = []

        try:
            for table_idx, table in enumerate(document.tables):
                table_rows: List[List[str]] = []

                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_cells)

                tables_data.append(
                    {
                        "table_index": table_idx,
                        "row_count": len(table.rows),
                        "col_count": len(table.columns) if table.rows else 0,
                        "data": table_rows,
                    }
                )

        except Exception as e:
            logger.warning(
                f"Failed to extract tables from {self.filename}: {str(e)}"
            )

        return tables_data

    def _extract_lists(self, document: Any) -> List[Dict[str, Any]]:
        """Extract list items from document.

        Args:
            document: python-docx Document object

        Returns:
            List of dictionaries containing list data
        """
        lists: List[Dict[str, Any]] = []
        current_list: List[str] = []
        in_list = False

        try:
            for para in document.paragraphs:
                # Check if paragraph is part of a list
                if para.style.name.startswith("List"):
                    if not in_list:
                        in_list = True
                        current_list = []
                    current_list.append(para.text.strip())
                else:
                    if in_list and current_list:
                        lists.append(
                            {
                                "type": "list",
                                "item_count": len(current_list),
                                "items": current_list,
                            }
                        )
                        current_list = []
                        in_list = False

            # Add the last list if exists
            if in_list and current_list:
                lists.append(
                    {
                        "type": "list",
                        "item_count": len(current_list),
                        "items": current_list,
                    }
                )

        except Exception as e:
            logger.warning(f"Failed to extract lists from {self.filename}: {str(e)}")

        return lists

    def _extract_formatting(self, document: Any) -> Dict[str, Any]:
        """Extract formatting information from document.

        Args:
            document: python-docx Document object

        Returns:
            Dict containing formatting statistics
        """
        formatting: Dict[str, Any] = {
            "paragraph_count": 0,
            "has_bold": False,
            "has_italic": False,
            "has_underline": False,
        }

        try:
            formatting["paragraph_count"] = len(document.paragraphs)

            # Check for formatting in runs
            for para in document.paragraphs:
                for run in para.runs:
                    if run.bold:
                        formatting["has_bold"] = True
                    if run.italic:
                        formatting["has_italic"] = True
                    if run.underline:
                        formatting["has_underline"] = True

                    # Early exit if all formatting types found
                    if (
                        formatting["has_bold"]
                        and formatting["has_italic"]
                        and formatting["has_underline"]
                    ):
                        break

        except Exception as e:
            logger.warning(
                f"Failed to extract formatting from {self.filename}: {str(e)}"
            )

        return formatting
